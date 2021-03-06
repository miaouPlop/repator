# coding=utf-8


import re
import json
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from conf.report import *
from src.cvss import *


class Generator:
    def __escape_str(s):
        return s.replace("\a", "\\a").replace("\b", "\\b").replace("\f", "\\f").replace("\r", "\\r").replace("\t",
                                                                                                             "\\t").replace(
            "\v", "\\v")

    def generate_report(json_content):
        if type(json_content) is str:
            return json_content.replace("%", "\%")

        if "name" not in json_content:
            json_content["name"] = ""
        if "content" not in json_content:
            json_content["content"] = ""

        file_content = None
        with open("templates/" + json_content["type"] + ".tex") as f:
            file_content = f.read()
        final_content = file_content
        final_content = re.sub("##.*NAME##", json_content["name"], final_content)

        content = ""
        if type(json_content["content"]) is list:
            for e in json_content["content"]:
                content += Generator.generate_report(e)
        elif type(json_content["content"]) is dict:
            content += Generator.generate_report(json_content["content"])
        elif type(json_content["content"]) is str:
            content += Generator.generate_report(json_content["content"])

        final_content = re.sub("##.*CONTENT##", content, final_content)
        return Generator.__escape_str(final_content)

    def __sub_dict(d, text):
        for name, value in d.items():
            if type(value) is str:
                text = re.sub("##" + name + "##", value, text)
                m = re.search("##" + name + r"#(\d+)##", text, re.MULTILINE)
                if m is not None:
                    text = re.sub("##" + name + "#" + m.group(1) + "##", value[int(m.group(1))], text)
                m = re.search("##" + name + r"#(\d+):(\d+)##", text, re.MULTILINE)
                if m is not None:
                    text = re.sub("##" + name + "#" + m.group(1) + ":" + m.group(2) + "##",
                                  value[int(m.group(1)):int(m.group(2))], text)
            else:
                text = re.sub("##" + name + "##", str(value), text)

        return text

    def __do_fill(d, content):
        if type(d) is str:
            d = Generator.__sub_dict(content, d)
            return d

        if type(d) is list:
            l = []
            for e in d:
                l.append(Generator.__do_fill(e, content))
            return l

        if "type" not in d:
            return d

        d["type"] = Generator.__sub_dict(content, d["type"])

        if d["type"] == "unordered_list" or d["type"] == "ordered_list":
            l = []
            for e in content[d["filer"]].values():
                for content in d["content"]:
                    template = dict(content)
                    template["content"] = Generator.__do_fill(template["content"], e)
                    l.append(template)
            d["content"] = l
            return d

        if d["type"].startswith("Vulns-"):
            l = []
            match = d["type"].split("-")[1]
            for docId, vuln in content["Vulns"].items():
                vuln["doc_id"] = docId
                if "riskLvl" not in vuln:
                    vuln["riskLvl"], vuln["impLvl"], vuln["expLvl"] = vulnRiskLevel(vuln)
                    vuln["cvss"], vuln["cvssImp"], vuln["cvssExp"] = vulnCvssv3(vuln);

                if (vuln["riskLvl"] == match and vuln["status"] == "Vulnerable") or vuln["status"] == match:
                    for content in d["content"]:
                        template = dict(content)
                        template["content"] = Generator.__do_fill(template["content"], vuln)
                        l.append(template)
            d["content"] = l
            return d

        if d["type"] == "VulnsFull":
            l = []
            cat = None
            sub_cat = None
            for docId, vuln in content["Vulns"].items():
                vuln["doc_id"] = docId
                if vuln["category"] != cat:
                    cat = vuln["category"]
                    sub_cat = None
                    template = Generator.__do_fill(dict(d["catContent"]), vuln)
                    l.append(template)

                if vuln["sub_category"] != sub_cat:
                    sub_cat = vuln["sub_category"]
                    template = Generator.__do_fill(dict(d["subcatContent"]), vuln)
                    l.append(template)

                for content in d["content-" + vuln["status"]]:
                    template = Generator.__do_fill(dict(content), vuln)
                    l.append(template)
            d["content"] = l
            return d

        if "filer" in d:
            content = content[d["filer"]]

        if "content" not in d:
            return d

        if type(d["content"]) is list:
            l = []
            for e in d["content"]:
                l.append(Generator.__do_fill(e, content))
            d["content"] = l
        elif type(d["content"]) is str:
            d["content"] = Generator.__sub_dict(content, d["content"])
        else:
            d["content"] = Generator.__do_fill(d["content"], content)

        return d

    def generate_json(json_content, template):
        template_path = REPORT_TEMPLATE_DIR + template + "/"
        structure = None
        with open(template_path + REPORT_TEMPLATE_MAIN, "r") as f:
            structure = f.read()
            structure = json.loads(structure)

        result_json = {}
        result_json["type"] = "report"
        result_json["content"] = []
        for e in structure:
            with open(template_path + e + ".json", "r") as f:
                file_content = f.read()
            json_file_content = json.loads(file_content)
            res = Generator.__do_fill(json_file_content, json_content)
            result_json["content"].append(res)
        return result_json

    def __cut_before(string, match):
        return string[string.find(match) + len(match):]

    def __cut_after(string, match):
        return string[0:string.find(match)]

    def __get_body(xml_content):
        out = Generator.__cut_before(xml_content, "<w:body>")
        out = Generator.__cut_after(out, "</w:body>")
        return out

    def __align(alignment):
        return getattr(WD_ALIGN_PARAGRAPH, alignment)

    def __vAlign(alignment):
        return getattr(WD_ALIGN_VERTICAL, alignment)

    def __generate_table(document, json):
        table = document.add_table(json["row"], json["col"], json["style"])
        for row in range(0, json["row"]):
            for col in range(0, json["col"]):
                Generator.generate_docx(table.cell(row, col),
                                        json["content"][row][col])
                if "alignment" in json:
                    for p in table.cell(row, col).paragraphs:
                        p.alignment = Generator.__align(json["alignment"])

                if "colAlignment" in json:
                    for p in table.cell(row, col).paragraphs:
                        p.alignment = Generator.__align(json["colAlignment"][col])
                if "rowAlignment" in json:
                    for p in table.cell(row, col).paragraphs:
                        p.alignment = Generator.__align(json["rowAlignment"][row])

                if "celAlignment" in json:
                    for p in table.cell(row, col).paragraphs:
                        p.alignment = Generator.__align(json["celAlignment"][row][col])

                if "vAlignment" in json:
                    table.cell(row, col).vertical_alignment = Generator.__vAlign(json["vAlignment"])

                if "colVAlignment" in json:
                    table.cell(row, col).vertical_alignment = Generator.__vAlign(json["colVAlignment"][col])

                if "rowVAlignment" in json:
                    table.cell(row, col).vertical_alignment = Generator.__vAlign(json["rowVAlignment"][row])

                if "celVAlignment" in json:
                    table.cell(row, col).vertical_alignment = Generator.__vAlign(json["celVAlignment"][row][col])

                if "width" in json:
                    col = 0
                    for width in json["width"]:
                        table.columns[col].width = Cm(width)
                        col += 1

                if "firstCol" in json:
                    table.first_col = json["firstCol"]
                if "firstRow" in json:
                    table.first_row = json["firstRow"]
                if "lastCol" in json:
                    table.last_col = json["lastCol"]
                if "lastRow" in json:
                    table.last_row = json["lastRow"]
                if "hBand" in json:
                    table.h_band(json["hBand"])
                if "vBand" in json:
                    table.v_band(json["vBand"])

    def generate_docx(document, json):
        if isinstance(json, str):
            document.text = json
            return

        if "type" in json:
            if json["type"] == "table":
                Generator.__generate_table(document, json)

            if json["type"] == "document":
                newDoc = Document(json["path"])

                for paragraph in newDoc.paragraphs:
                    text = paragraph.text
                    style = paragraph.style.name
                    p = document.add_paragraph(text, style)
                    p.paragraph_format = paragraph.paragraph_format

        if "name" in json:
            p = document.add_paragraph(json["name"], json["type"])
            if "alignment" in json:
                p.alignment = Generator.__align(json["alignment"])

        if "content" in json:
            if isinstance(json["content"], str):
                contentTable = json["content"].split("[[HYPERLINK]]")
                if "type" not in json:
                    json["type"] = "Normal"

                p = document.add_paragraph(contentTable[0], json["type"])

                if "alignment" in json:
                    p.alignment = Generator.__align(json["alignment"])

                for cpt in range(1, len(contentTable)):
                    if cpt % 2:
                        hyperlink = contentTable[cpt].split("||")
                        p.add_hyperlink(hyperlink[0], hyperlink[1], style="Hyperlink")
                    else:
                        p.add_run(contentTable[cpt])

            elif isinstance(json["content"], list):
                for content in json["content"]:
                    Generator.generate_docx(document, content)
            else:
                Generator.generate_docx(document, json["content"])

    def generate_all(values, outputFilename):
        template = values["Mission"]["template"]

        p = Generator.generate_json(values, template)

        doc = Document(docx=REPORT_TEMPLATE_DIR + template + "/" + REPORT_TEMPLATE_BASE)
        Generator.generate_docx(doc, p)
        doc.save(outputFilename)

### testing
# d = {"general": {"date_start": "11/11/11",
#                  "date_end": "12/12/12"},

#      "clients": [{"name": "john doe",
#                   "email": "john@doe.com",
#                   "tel": "+33 66 55 44 33",
#                   "role": "CEO"}],

#      "auditors": [{"name": "haxor auditor",
#                    "email": "haxor@auditor.com",
#                    "tel": "31337",
#                    "role": "pro-hacker"},

#                   {"name": "script kiddie",
#                    "email": "skiddie@mail.com",
#                    "tel": "123456",
#                    "role": "skiddie"}]}

# p = (Generator.generate_json(d))
# #print(p)
# #print(Generator.generate_report(p))
# doc = Document(docx="templates/template.docx")
# Generator.generate_docx(doc, p)
# doc.save("test.docx")
